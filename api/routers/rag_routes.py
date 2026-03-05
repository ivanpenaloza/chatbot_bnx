"""
RAG Router — Satriani Document Chat

Provides REST API endpoints for:
  1. Ingesting .docx/.pdf files → extract text → chunk → embed → ChromaDB
  2. Querying ChromaDB for relevant chunks
  3. Unified chat: optional RAG context + optional uploaded doc context
  4. User document upload (temporary, for chat context)

Compatible with Python 3.9.20.
"""

import os
import sys
import re
import logging
import hashlib
import traceback
import tempfile
from typing import Optional, List, Dict, Any

from fastapi import APIRouter, HTTPException, UploadFile, File, Request
from fastapi.responses import JSONResponse, HTMLResponse
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel

project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

from config import (
    MODELS_BASE_DIR,
    RAG_DATA_DIR,
    RAG_UPLOAD_MAX_SIZE_MB,
    AVAILABLE_EMBEDDING_MODELS,
    DEFAULT_EMBEDDING_KEY,
    CHROMA_PERSIST_DIR,
)

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/v1", tags=["rag"])
templates = Jinja2Templates(directory="templates")

# ─── Globals (lazy-loaded) ───────────────────────────────────────────────────
_chroma_client = None
_embedding_fns: Dict[str, Any] = {}  # cache per embedding key


# ─── Pydantic Models ─────────────────────────────────────────────────────────

class RagChatMessage(BaseModel):
    message: str
    collections: Optional[List[str]] = None
    use_rag: Optional[bool] = True
    uploaded_context: Optional[List[Dict[str, str]]] = None  # [{filename, text}]
    chat_history: Optional[List[Dict[str, str]]] = None  # [{role, content}]


class IngestRequest(BaseModel):
    filename: Optional[str] = None
    embedding_key: Optional[str] = None


# ─── Embedding Function ──────────────────────────────────────────────────────

class LocalHFEmbeddingFunction:
    """ChromaDB-compatible embedding function using local sentence-transformers."""

    def __init__(self, model_path: str, model_key: str):
        from sentence_transformers import SentenceTransformer
        import torch
        import numpy as np
        device = "cuda" if torch.cuda.is_available() else "cpu"
        self._model = SentenceTransformer(model_path, trust_remote_code=True, device=device)
        self._name = "local-hf-" + model_key
        self._np = np
        logger.info("Embedding model loaded on %s from: %s", device.upper(), model_path)


    def name(self) -> str:
        return self._name

    def __call__(self, input: List[str]) -> List:
        embeddings = self._model.encode(input, show_progress_bar=False)
        return [self._np.array(e, dtype=self._np.float32) for e in embeddings]

    def embed_query(self, input: List[str]) -> List:
        return self.__call__(input)

    def embed_with_retries(self, input, **kwargs):
        return self.__call__(input)

    def default_space(self) -> str:
        return "cosine"

    def supported_spaces(self):
        return ["cosine", "l2", "ip"]

    def get_config(self):
        return {"model_path": self._name}

    @staticmethod
    def validate_config(config):
        pass

    def validate_config_update(self, old_config, new_config):
        pass

    @classmethod
    def build_from_config(cls, config):
        raise NotImplementedError

    def is_legacy(self) -> bool:
        return False


def get_embedding_fn(embedding_key: Optional[str] = None):
    """Get or create an embedding function for the given key. Cached per key."""
    key = embedding_key or DEFAULT_EMBEDDING_KEY
    if key not in _embedding_fns:
        model_info = AVAILABLE_EMBEDDING_MODELS.get(key)
        if not model_info:
            raise ValueError(f"Unknown embedding model: {key}")
        _embedding_fns[key] = LocalHFEmbeddingFunction(model_info["local_dir"], key)
    return _embedding_fns[key]


def get_chroma_client():
    global _chroma_client
    if _chroma_client is None:
        import chromadb
        from chromadb.config import Settings
        os.makedirs(CHROMA_PERSIST_DIR, exist_ok=True)
        _chroma_client = chromadb.PersistentClient(
            path=CHROMA_PERSIST_DIR,
            settings=Settings(anonymized_telemetry=False),
        )
        logger.info("ChromaDB initialized at: %s", CHROMA_PERSIST_DIR)
    return _chroma_client


# ─── Text Extraction ─────────────────────────────────────────────────────────

def extract_text_from_docx(filepath: str) -> str:
    from docx import Document
    doc = Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)






def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from PDF using pymupdf (fitz).

    pymupdf uses the MuPDF engine which preserves word boundaries
    from the PDF layout, producing clean readable text.
    """
    import fitz
    text_parts = []
    with fitz.open(filepath) as doc:
        for page in doc:
            t = page.get_text("text")
            if t:
                text_parts.append(t.strip())
    raw = "\n".join(text_parts)
    return _fix_pdf_spacing(raw)






def extract_text(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".docx":
        return extract_text_from_docx(filepath)
    elif ext == ".pdf":
        return extract_text_from_pdf(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ─── Chunking ────────────────────────────────────────────────────────────────


def _clean_extracted_text(text: str) -> str:
    """Clean up common PDF/docx extraction artifacts for display.

    Handles concatenated words (e.g. 'theenzymesare' → 'the enzymes are')
    by detecting likely word boundaries using case transitions, common
    patterns, and heuristics.
    """
    # First pass: fix spacing issues from PDF extraction
    text = _fix_pdf_spacing(text)
    # Collapse multiple whitespace into single space
    text = re.sub(r'[ \t]+', ' ', text)
    # Collapse 3+ newlines into 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def _fix_pdf_spacing(text: str) -> str:
    """Re-insert spaces into text where PDF extraction concatenated words.

    Uses multiple heuristics:
      1. Insert space between a lowercase letter and an uppercase letter
         (e.g. 'enzymesAre' → 'enzymes Are')
      2. Insert space between a letter and a digit boundary
         (e.g. 'pH7' → 'pH 7', 'level5' → 'level 5')
      3. Insert space between a period/comma and a letter with no space
         (e.g. 'cells.The' → 'cells. The', 'acids,which' → 'acids, which')
      4. Insert space between a closing paren and a letter
         (e.g. ')The' → ') The')
    """
    if not text:
        return text
    # lowercase followed by uppercase: camelCase word boundaries
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
    # letter followed by digit (but not common patterns like pH, H2O, etc.)
    text = re.sub(r'([a-zA-Z]{2,})(\d)', r'\1 \2', text)
    # digit followed by letter (e.g. '5in' → '5 in', but keep '2nd', '3rd', etc.)
    text = re.sub(r'(\d)([a-zA-Z])', lambda m: m.group(0) if m.group(2) in ('st','nd','rd','th') else m.group(1) + ' ' + m.group(2), text)
    # period/comma/semicolon/colon followed by a letter (no space)
    text = re.sub(r'([.,;:])([A-Za-z])', r'\1 \2', text)
    # closing paren followed by a letter
    text = re.sub(r'\)([A-Za-z])', r') \1', text)
    # letter followed by opening paren
    text = re.sub(r'([a-zA-Z])\(', r'\1 (', text)
    return text




def chunk_text(text: str, chunk_size: int = 500, overlap: int = 100) -> List[str]:
    """Split text into chunks, breaking at sentence boundaries when possible."""
    text = text.strip()
    if not text:
        return []

    # Split into sentences first for cleaner boundaries
    sentences = re.split(r'(?<=[.!?])\s+', text)

    chunks = []
    current_chunk = ""

    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        # If adding this sentence exceeds chunk_size, save current and start new
        if current_chunk and len(current_chunk) + len(sentence) + 1 > chunk_size:
            chunks.append(current_chunk.strip())
            # Keep overlap: take the last `overlap` characters of current chunk
            if overlap > 0 and len(current_chunk) > overlap:
                current_chunk = current_chunk[-overlap:].lstrip() + " " + sentence
            else:
                current_chunk = sentence
        else:
            current_chunk = (current_chunk + " " + sentence).strip() if current_chunk else sentence

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return chunks


def sanitize_collection_name(name: str, embedding_key: Optional[str] = None) -> str:
    """Build a ChromaDB collection name from filename + embedding model key.

    Including the model key allows the same document to be embedded with
    different models while preventing duplicate ingestion for the same
    (document, model) pair.
    """
    base = os.path.splitext(name)[0]
    sanitized = re.sub(r'[^a-zA-Z0-9_-]', '_', base)
    sanitized = sanitized.strip('_-')

    # Append a short model suffix so each model gets its own collection
    if embedding_key:
        model_suffix = re.sub(r'[^a-zA-Z0-9_-]', '_', embedding_key)
        sanitized = f"{sanitized}__{model_suffix}"

    if len(sanitized) < 3:
        sanitized = sanitized + "_doc"
    if len(sanitized) > 63:
        sanitized = sanitized[:63].rstrip('_-')
    return sanitized


# ─── Ingestion Logic ─────────────────────────────────────────────────────────

def ingest_file(filepath: str, embedding_key: Optional[str] = None) -> Dict[str, Any]:
    """Ingest a single docx/pdf file into ChromaDB.

    The collection name encodes both the filename and the embedding model,
    so the same document can be embedded with different models but cannot
    be ingested twice with the same model.
    """
    filename = os.path.basename(filepath)
    emb_key = embedding_key or DEFAULT_EMBEDDING_KEY
    collection_name = sanitize_collection_name(filename, emb_key)

    logger.info("Ingesting: %s → collection: %s (model: %s)", filename, collection_name, emb_key)

    # ── Duplicate check: same document + same model ──
    client = get_chroma_client()
    try:
        existing = client.get_collection(name=collection_name)
        if existing.count() > 0:
            logger.info(
                "Collection '%s' already has %d chunks (model: %s), skipping duplicate.",
                collection_name, existing.count(), emb_key,
            )
            return {
                "filename": filename,
                "collection": collection_name,
                "status": "already_ingested",
                "chunks": existing.count(),
                "embedding_model": emb_key,
            }
    except Exception:
        pass  # Collection doesn't exist yet — proceed

    text = extract_text(filepath)
    if not text.strip():
        return {"filename": filename, "status": "empty", "chunks": 0}

    chunks = chunk_text(text, chunk_size=500, overlap=100)

    ef = get_embedding_fn(emb_key)

    collection = client.get_or_create_collection(
        name=collection_name,
        embedding_function=ef,
        metadata={"source_file": filename, "embedding_model": emb_key},
    )

    ids = []
    for i, chunk in enumerate(chunks):
        chunk_hash = hashlib.md5(chunk.encode()).hexdigest()[:8]
        ids.append(f"{collection_name}_{i}_{chunk_hash}")

    metadatas = [{"source": filename, "chunk_index": i} for i in range(len(chunks))]

    batch_size = 100
    for start in range(0, len(chunks), batch_size):
        end = min(start + batch_size, len(chunks))
        collection.add(
            ids=ids[start:end],
            documents=chunks[start:end],
            metadatas=metadatas[start:end],
        )

    logger.info("Ingested %d chunks for '%s'", len(chunks), filename)
    return {
        "filename": filename,
        "collection": collection_name,
        "status": "ingested",
        "chunks": len(chunks),
        "text_length": len(text),
        "embedding_model": emb_key,
    }


def get_all_rag_files() -> List[str]:
    """Return list of .docx and .pdf files in RAG_DATA_DIR."""
    files = []
    if os.path.isdir(RAG_DATA_DIR):
        for f in sorted(os.listdir(RAG_DATA_DIR)):
            if f.lower().endswith(('.docx', '.pdf')):
                files.append(os.path.join(RAG_DATA_DIR, f))
    return files


# ─── Query Logic ──────────────────────────────────────────────────────────────

def query_collections(
    question: str,
    collection_names: Optional[List[str]] = None,
    n_results: int = 3,
) -> List[Dict[str, Any]]:
    client = get_chroma_client()
    ef = get_embedding_fn()

    all_collections = client.list_collections()
    if collection_names:
        target_names = collection_names
    else:
        target_names = [c.name for c in all_collections]

    if not target_names:
        return []

    results = []
    for cname in target_names:
        try:
            collection = client.get_collection(name=cname, embedding_function=ef)
            if collection.count() == 0:
                continue
            query_result = collection.query(
                query_texts=[question],
                n_results=min(n_results, collection.count()),
                include=["documents", "metadatas", "distances"],
            )
            docs = query_result.get("documents", [[]])[0]
            metas = query_result.get("metadatas", [[]])[0]
            dists = query_result.get("distances", [[]])[0]
            col_meta = collection.metadata or {}
            emb_model = col_meta.get("embedding_model", "unknown")
            for doc, meta, dist in zip(docs, metas, dists):
                results.append({
                    "collection": cname,
                    "document": doc,
                    "source": meta.get("source", ""),
                    "chunk_index": meta.get("chunk_index", 0),
                    "distance": round(dist, 4),
                    "embedding_model": emb_model,
                })
        except Exception as e:
            logger.warning("Error querying collection '%s': %s", cname, e)

    results.sort(key=lambda x: x["distance"])
    return results[:n_results]


# ─── Routes ───────────────────────────────────────────────────────────────────

@router.get("/rag/chat", response_class=HTMLResponse)
async def rag_chat_page(request: Request):
    return templates.TemplateResponse("rag_chat.html", {"request": request})


@router.post("/rag/ingest", response_class=JSONResponse)
async def rag_ingest(req: IngestRequest):
    """Ingest a SINGLE file into ChromaDB. filename is required."""
    try:
        if not req.filename:
            raise HTTPException(status_code=400, detail="filename is required")

        files = get_all_rag_files()
        matched = [f for f in files if os.path.basename(f) == req.filename]
        if not matched:
            raise HTTPException(status_code=404, detail=f"File not found: {req.filename}")

        result = ingest_file(matched[0], req.embedding_key)
        return JSONResponse(content={"status": "ok", "results": [result]})
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Ingestion error: %s", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/rag/documents")
async def rag_list_documents():
    """List available documents and their ingestion status from ChromaDB.

    A document may have multiple embeddings (one per model), so we return
    a list of embedding entries per file.
    """
    files = get_all_rag_files()
    client = get_chroma_client()

    # Index all collections by source_file from metadata
    collections_by_source: Dict[str, List[Any]] = {}
    for col in client.list_collections():
        meta = col.metadata or {}
        src = meta.get("source_file", "")
        if src:
            collections_by_source.setdefault(src, []).append(col)

    docs = []
    for fpath in files:
        fname = os.path.basename(fpath)
        matched_cols = collections_by_source.get(fname, [])
        if matched_cols:
            for col in matched_cols:
                col_meta = col.metadata or {}
                emb_model = col_meta.get("embedding_model", "")
                emb_algorithm = col_meta.get("hnsw:space", "cosine")
                try:
                    chunk_count = col.count()
                except Exception:
                    chunk_count = 0
                docs.append({
                    "filename": fname,
                    "collection": col.name,
                    "ingested": True,
                    "chunks": chunk_count,
                    "size_kb": round(os.path.getsize(fpath) / 1024, 1),
                    "embedding_model": emb_model,
                    "embedding_algorithm": emb_algorithm or "cosine",
                })
        else:
            # Document exists on disk but has no embeddings yet
            docs.append({
                "filename": fname,
                "collection": sanitize_collection_name(fname),
                "ingested": False,
                "chunks": 0,
                "size_kb": round(os.path.getsize(fpath) / 1024, 1),
                "embedding_model": "",
                "embedding_algorithm": "cosine",
            })
    return JSONResponse(content={"documents": docs})


@router.post("/rag/upload-context")
async def upload_context_document(file: UploadFile = File(...)):
    """User uploads a doc/pdf for temporary chat context (max 5MB).
    Extracts text and returns it — does NOT persist to ChromaDB."""
    fname = file.filename or "unknown"
    ext = os.path.splitext(fname)[1].lower()
    if ext not in (".docx", ".pdf"):
        raise HTTPException(status_code=400, detail="Only .docx and .pdf files are allowed")

    content = await file.read()
    size_mb = len(content) / (1024 * 1024)
    if size_mb > RAG_UPLOAD_MAX_SIZE_MB:
        raise HTTPException(
            status_code=400,
            detail=f"File too large ({size_mb:.1f} MB). Max is {RAG_UPLOAD_MAX_SIZE_MB} MB."
        )

    suffix = ext
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        text = extract_text(tmp_path)
    finally:
        os.unlink(tmp_path)

    if not text.strip():
        return JSONResponse(content={"filename": fname, "text": "", "chunks": 0})

    chunks = chunk_text(text, chunk_size=500, overlap=100)
    return JSONResponse(content={
        "filename": fname,
        "text": text[:50000],
        "chunks": len(chunks),
        "size_mb": round(size_mb, 2),
    })


@router.post("/rag/ask")
async def rag_ask(chat: RagChatMessage):
    """Unified chat with intelligent flow routing.

    Flow graph:
      ┌─────────────┐
      │  User Query  │
      └──────┬───────┘
             │
      ┌──────▼───────┐
      │ Classify      │──→ has_rag? has_upload? neither?
      └──────┬───────┘
             │
      ┌──────▼───────────────────────────────────────┐
      │  BRANCH A: RAG + Upload (merge & rank)       │
      │  BRANCH B: RAG only                          │
      │  BRANCH C: Upload only                       │
      │  BRANCH D: No docs → identity prompt         │
      └──────┬───────────────────────────────────────┘
             │
      ┌──────▼───────┐
      │ Build prompt  │
      │ + LLM call    │
      └──────────────┘
    """
    if not chat.message.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty")

    rag_sources = []
    uploaded_sources = []
    rag_context = ""
    uploaded_context = ""

    # ── Step 1: RAG retrieval ────────────────────────────────────────────
    has_rag = False
    if chat.use_rag and chat.collections:
        chunks = query_collections(
            question=chat.message,
            collection_names=chat.collections,
            n_results=3,
        )
        if chunks:
            # Filter by relevance threshold — discard chunks with very low relevance
            RELEVANCE_THRESHOLD = 0.25  # minimum relevance (1 - distance)
            relevant_chunks = [c for c in chunks if (1 - c['distance']) >= RELEVANCE_THRESHOLD]

            if relevant_chunks:
                has_rag = True
                context_parts = []
                for i, chunk in enumerate(relevant_chunks, 1):
                    cleaned = _clean_extracted_text(chunk['document'])
                    context_parts.append(
                        f"[Source {i}: {chunk['source']}, "
                        f"relevance: {1 - chunk['distance']:.0%}]\n"
                        f"{cleaned}"
                    )
                    rag_sources.append({
                        "type": "rag",
                        "source": chunk["source"],
                        "collection": chunk["collection"],
                        "chunk_index": chunk["chunk_index"],
                        "relevance": round(1 - chunk["distance"], 4),
                        "text": cleaned,
                        "embedding_model": chunk.get("embedding_model", ""),
                    })
                rag_context = "\n\n".join(context_parts)
            else:
                # Chunks found but all below relevance threshold — still use
                # document prompt so the model can say "not related"
                has_rag = True
                rag_context = (
                    "[NOTE: The retrieved document fragments have very low relevance "
                    "to the user's question. The question is likely unrelated to the "
                    "loaded documents. Inform the user accordingly.]"
                )

    # ── Step 2: Uploaded document processing (top 3 documents max) ──────
    has_upload = False
    if chat.uploaded_context:
        ctx_parts = []
        for i, doc in enumerate(chat.uploaded_context[:3], 1):
            fname = doc.get("filename", f"uploaded_{i}")
            text = doc.get("text", "")
            if text.strip():
                cleaned = _clean_extracted_text(text)
                doc_chunks = chunk_text(cleaned, chunk_size=600, overlap=50)
                selected_text = ""
                for dc in doc_chunks:
                    if len(selected_text) + len(dc) > 5000:
                        break
                    selected_text += dc + "\n\n"
                selected_text = selected_text.strip()
                if selected_text:
                    has_upload = True
                    ctx_parts.append(
                        f"[Uploaded: {fname}]\n{selected_text}"
                    )
                    preview = selected_text[:500] + ("..." if len(selected_text) > 500 else "")
                    uploaded_sources.append({
                        "type": "uploaded",
                        "source": fname,
                        "collection": "user_upload",
                        "chunk_index": 0,
                        "relevance": 1.0,
                        "text": preview,
                        "embedding_model": "n/a (raw text)",
                    })
        if ctx_parts:
            uploaded_context = "\n\n".join(ctx_parts)

    # ── Step 3: Route to the correct branch ──────────────────────────────
    from config import SATRIANI_IDENTITY_PROMPT, SATRIANI_DOCUMENT_PROMPT

    if has_rag and has_upload:
        # BRANCH A: Both RAG and uploaded docs — merge and rank
        system_prompt = SATRIANI_DOCUMENT_PROMPT
        data_context = (
            "You have access to TWO types of document sources. "
            "Cross-reference them to give the most complete answer.\n\n"
            "KNOWLEDGE BASE DOCUMENTS (from RAG):\n" + rag_context + "\n\n"
            "UPLOADED DOCUMENTS (user-provided):\n" + uploaded_context
        )
    elif has_rag:
        # BRANCH B: RAG only
        system_prompt = SATRIANI_DOCUMENT_PROMPT
        data_context = (
            "KNOWLEDGE BASE DOCUMENTS:\n" + rag_context
        )
    elif has_upload:
        # BRANCH C: Uploaded docs only
        system_prompt = SATRIANI_DOCUMENT_PROMPT
        data_context = (
            "UPLOADED DOCUMENTS:\n" + uploaded_context
        )
    else:
        # BRANCH D: No documents — use identity prompt, general assistant
        system_prompt = SATRIANI_IDENTITY_PROMPT
        data_context = ""

    # ── Step 4: Build conversation history for proper multi-turn chat ────
    conversation_history = []
    if chat.chat_history and len(chat.chat_history) > 0:
        for msg in chat.chat_history[-6:]:
            role = msg.get("role", "user")
            content = msg.get("content", "")
            if content:
                conversation_history.append({
                    "role": role if role in ("user", "assistant") else "user",
                    "content": content[:1000],
                })

    # ── Step 5: Generate response ────────────────────────────────────────
    from .llm_routes import model_manager
    answer, inference_time = model_manager.generate_response(
        user_message=chat.message,
        data_context=data_context,
        system_prompt_override=system_prompt,
        conversation_history=conversation_history,
    )

    return JSONResponse(content={
        "response": answer,
        "sources": rag_sources + uploaded_sources,
        "rag_sources": rag_sources,
        "uploaded_sources": uploaded_sources,
        "inference_time": inference_time,
        "chunks_used": len(rag_sources) + len(uploaded_sources),
    })
